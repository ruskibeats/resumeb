"""
Generate CV_Russell_Batchelor_Master_FINAL.docx from master markdown content.
Uses python-docx with professional UK CV formatting.
2-page optimised version — tight spacing, merged concurrent roles, trimmed bullets.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x1A, 0x2E, 0x44)   # headings
MID_GREY    = RGBColor(0x55, 0x5F, 0x6E)   # sub-headings / meta
RULE_COLOUR = RGBColor(0x1A, 0x2E, 0x44)   # horizontal rules
BODY_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)   # body text

# ── Helper: add horizontal rule ───────────────────────────────────────────────
def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A2E44')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Helper: set paragraph spacing ─────────────────────────────────────────────
def spacing(para, before=0, after=2, line=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if line:
        para.paragraph_format.line_spacing = line

# ── Helper: bold run ──────────────────────────────────────────────────────────
def bold_run(para, text, size=10, colour=BODY_BLACK):
    r = para.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = colour
    return r

# ── Helper: normal run ────────────────────────────────────────────────────────
def normal_run(para, text, size=10, colour=BODY_BLACK, italic=False):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = colour
    r.italic = italic
    return r

# ══════════════════════════════════════════════════════════════════════════════
# NAME / CONTACT HEADER
# ══════════════════════════════════════════════════════════════════════════════
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(name_para, before=0, after=2)
r = name_para.add_run("Russell Batchelor")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = DARK_NAVY

contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(contact_para, before=0, after=2)
normal_run(contact_para,
    "Hemel Hempstead, Hertfordshire  |  linkedin.com/in/russellbatchelor  |  "
    "SC Cleared (Active)  |  NPPV3 (Active)",
    size=9, colour=MID_GREY)

avail_para = doc.add_paragraph()
avail_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(avail_para, before=0, after=4)
normal_run(avail_para,
    "Available within 2 weeks  |  UK-based; internationally mobile  |  "
    "Hiloka Limited  |  £650/day",
    size=9, colour=MID_GREY)

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# HEADLINE
# ══════════════════════════════════════════════════════════════════════════════
hl = doc.add_paragraph()
hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(hl, before=4, after=2)
r = hl.add_run("Infrastructure Programme Director  |  Technical Operations Director")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = DARK_NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(sub, before=0, after=4)
normal_run(sub,
    "Network, Data Centre & Service Transformation  ·  Managed Services  ·  "
    "Regulated & Mission-Critical Estates  ·  UK & EMEA",
    size=9, colour=MID_GREY, italic=True)

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PROFILE
# ══════════════════════════════════════════════════════════════════════════════
profile_text = (
    "Senior infrastructure and technical operations programme director with 30 years delivering "
    "high-stakes transformation across regulated UK and EMEA environments. Brought in by boards, "
    "CEOs, and PE investors to mobilise stalled programmes and turn complex infrastructure estates "
    "into commercially governed operations. Combines delivery authority (discovery through cutover) "
    "with full P&L ownership, SLA governance, and board-level reporting. Proven across NATO, NHS, "
    "central government, Tier 1 financial services, and PE-backed technology businesses. "
    "SC and NPPV3 cleared (both active). Available within 2 weeks."
)
p = doc.add_paragraph()
spacing(p, before=3, after=4)
normal_run(p, profile_text, size=9.5)

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CORE SKILLS (3-column table)
# ══════════════════════════════════════════════════════════════════════════════
skills_heading = doc.add_paragraph()
spacing(skills_heading, before=4, after=3)
bold_run(skills_heading, "CORE SKILLS", size=10, colour=DARK_NAVY)

skills = [
    ["Infrastructure Transformation Delivery",  "Technical Operations Leadership",         "Network Modernisation (LAN/WAN/Wi-Fi, SD-WAN, SASE)"],
    ["Data Centre Transition & Service Readiness", "EUC / Workplace Technology Programmes", "Programme Governance, RAID & Assurance"],
    ["P&L Accountability & Cost-to-Serve",      "SLA & Vendor Performance Governance",    "Contract-to-Cash & Revenue Assurance"],
    ["Regulated / Mission-Critical Delivery",   "Executive & Board-Level Reporting",      "UK & EMEA Multi-Site Delivery"],
]

tbl = doc.add_table(rows=len(skills), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

for ri, row_data in enumerate(skills):
    for ci, cell_text in enumerate(row_data):
        cell = tbl.cell(ri, ci)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cell_text)
        r.font.size = Pt(8.5)
        r.font.color.rgb = BODY_BLACK
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        # light background on even rows
        if ri % 2 == 0:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EEF2F7')
            tc_pr.append(shd)

doc.add_paragraph()  # spacer

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SELECTED CAREER OUTCOMES
# ══════════════════════════════════════════════════════════════════════════════
co_heading = doc.add_paragraph()
spacing(co_heading, before=4, after=3)
bold_run(co_heading, "SELECTED CAREER OUTCOMES", size=10, colour=DARK_NAVY)

outcomes = [
    ("60% WAN cost reduction", " across 200+ Lambeth sites; led ", "£12M voice & data re-tender", " — full commercial lifecycle from strategy through award."),
    ("De-facto EMEA Operations Director", " (CentricsIT, 4 years) — ", "15+ concurrent programmes", "; NATO, Tier 1 finance, NHS, UK Government; ", "6,100+ circuits", ", ", "7,000+ endpoints", "."),
    ("Country Operations Director", " (Gaming & HFT Group, Bulgaria) — P&L and ", "24-person in-country team", "; ", "£3M low-latency network build", ", 1,200+ endpoints, 7 IDCs."),
    ("NATO pan-European network upgrade", " — ", "200+ airbases", ", ", "£10M+ programme", ", 30-person cross-border team, on-time under SC mandate."),
    ("Home Office Microsoft 365", " (~42,000 users) and 7 central government departments under GDS; ", "25,000 Wi-Fi endpoints", " and 7 DC transitions (Sitehands)."),
]

for parts in outcomes:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.2)
    spacing(p, before=1, after=2)
    bold_flag = True
    for part in parts:
        if bold_flag and part:
            bold_run(p, part, size=9.5)
        else:
            normal_run(p, part, size=9.5)
        bold_flag = not bold_flag

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════════════════════════════════════════
exp_heading = doc.add_paragraph()
spacing(exp_heading, before=4, after=3)
bold_run(exp_heading, "PROFESSIONAL EXPERIENCE", size=10, colour=DARK_NAVY)

# ── Role data ─────────────────────────────────────────────────────────────────
roles = [
    {
        "employer": "Atos  /  Collective IP",
        "title": "Infrastructure Programme Manager  |  Technical Programme Manager",
        "dates": "Jun 2025 – Present  |  London, South East & Remote  |  Contract via Hiloka Limited  (concurrent engagements)",
        "note":  None,
        "mandate": None,
        "bullets": [
            "Atos: Programme lead for Windows 11 and network modernisation across live NHS Trust estates — "
            "directing site readiness, change windows, and LAN/Wi-Fi uplifts under strict zero-disruption constraints.",
            "Collective IP: Senior delivery lead for concurrent EMEA infrastructure transformation programmes "
            "(network separation, EUC upgrades, DC transitions) across Tier 1 financial services and public sector.",
            "Both: established RAID governance, escalation frameworks, and executive reporting across complex "
            "multi-stakeholder environments.",
        ],
        "tech": "Technologies: LAN/WAN, Wi-Fi, EUC, data centre  |  Sectors: NHS, financial services, public sector",
    },
    {
        "employer": "Hiloka Limited",
        "title": "Technology, Risk & Programme Assurance (independent advisory)",
        "dates": "Feb 2025 – Sept 2025  |  UK / Remote",
        "note":  None,
        "mandate": None,
        "bullets": [
            "Discovery, assurance, and recovery advisory for high-risk or stalled transformation programmes — "
            "root-cause assessments, vendor capability reviews, and executive-ready reset recommendations.",
        ],
        "tech": None,
    },
    {
        "employer": "Major International Gaming & HFT Group",
        "title": "Operations Director / Programme Director",
        "dates": "Jan 2024 – Feb 2025  |  Bulgaria, EU  |  Contract via Hiloka Limited",
        "note":  None,
        "mandate": (
            "Country-level Operations Director — greenfield digital gaming and HFT platform for major "
            "international group's European market entry; full P&L accountability, no prior blueprint."
        ),
        "bullets": [
            "Owned full P&L and contract-to-cash lifecycle: vendor contracts, SLA frameworks, revenue assurance, "
            "and margin governance; 1,200+ endpoints, primary DC, and 7 IDCs in live operations.",
            "Led 24-person multilingual engineering team; chaired Board-level KPI packs covering performance, "
            "risk, compliance, and commercial position from Day 1.",
            "Designed and delivered end-to-end operational handover model — transitioning greenfield build into "
            "auditable, board-reportable live service.",
        ],
        "tech": "Technologies: Low-latency networking, data centre, AWS, SLA governance  |  Sectors: Digital gaming, HFT",
    },
    {
        "employer": "CentricsIT (now Park Place Technologies)",
        "title": "EMEA Operations & Infrastructure Director (de facto)",
        "dates": "Nov 2020 – Dec 2024  |  UK & Europe  |  Contract via Hiloka Limited",
        "note":  None,
        "mandate": (
            "De-facto EMEA Operations Director — 4-year tenure; 15+ concurrent programmes across "
            "Tier 1 finance, UK Government, NATO, and NHS."
        ),
        "bullets": [
            "6,100+ circuit migrations, 7,000+ endpoint deployments, 130+ site rollouts — SD-WAN/SASE, "
            "LAN/WAN/Wi-Fi, data centre, and EUC programmes across regulated EMEA estates.",
            "NATO pan-European network upgrade: 200+ airbases, £10M+ programme, 30-person cross-border team — "
            "flew to Germany to restructure on-ground delivery when programme was slipping; recovered on time.",
            "$5M+ SD-WAN programme (200+ US retail sites, 20 months) and full commercial lifecycle governance "
            "across all programmes: SLA performance, billing capture, vendor management, C-Suite reporting.",
        ],
        "tech": "Technologies: SD-WAN, SASE, LAN/WAN/Wi-Fi, data centre, EUC  |  Clients: NATO, Tier 1 finance, UK Government, NHS",
    },
    {
        "employer": "Sitehands LLC",
        "title": "Infrastructure Project Director, Europe",
        "dates": "Nov 2018 – Nov 2020  |  London & Europe  |  Contract via Hiloka Limited",
        "note":  None,
        "mandate": None,
        "bullets": [
            "25,000 Wi-Fi endpoints and 7 data centre transitions (on-prem to off-prem) across M&A-driven "
            "European estate transformation; modernised legacy networks to SD-WAN and SASE architectures.",
            "Built operational control frameworks for consistent multi-vendor, multi-country programme execution.",
        ],
        "tech": "Technologies: LAN/WAN/Wi-Fi, SD-WAN, SASE, data centre migration",
    },
    {
        "employer": "Rainmaker Solutions",
        "title": "Associate Technology Consultant",
        "dates": "Nov 2015 – Nov 2018  |  London  |  Contract via Hiloka Limited",
        "note":  None,
        "mandate": None,
        "bullets": [
            "Microsoft 365 and EUC transformations across 7 central government departments under GDS — "
            "including Home Office (~42,000 users) and MHCLG (~5,000 users).",
            "Managed procurement strategy, bid authoring, design assurance, and delivery coordination across "
            "networking, cloud, and virtualisation platforms.",
        ],
        "tech": "Technologies: Microsoft 365, EUC, cloud, enterprise networking, GDS  |  Departments: Home Office, MHCLG, MoJ, NHS",
    },
    {
        "employer": "London Borough of Lambeth",
        "title": "Lead Enterprise Architect  (+ Network Project Manager, 2009–2010)",
        "dates": "Mar 2009 – Apr 2015  |  London  |  Contract",
        "note":  None,
        "mandate": None,
        "bullets": [
            "60% WAN cost reduction across 200+ sites; led £12M voice and data services re-tender — "
            "full commercial lifecycle from strategy through award and transition.",
            "Deployed telephony for 5,000 employees across 140+ sites; de-facto technology operations "
            "lead retained over 7 years.",
        ],
        "tech": "Technologies: LAN/WAN, telephony, strategic procurement  |  Sector: UK local government",
    },
]

for role in roles:
    # Employer + title
    emp_p = doc.add_paragraph()
    spacing(emp_p, before=5, after=0)
    bold_run(emp_p, role["employer"], size=10, colour=DARK_NAVY)
    normal_run(emp_p, "  —  " + role["title"], size=9.5, colour=DARK_NAVY)

    # Dates
    dates_p = doc.add_paragraph()
    spacing(dates_p, before=0, after=1)
    normal_run(dates_p, role["dates"], size=8.5, colour=MID_GREY, italic=True)
    if role["note"]:
        normal_run(dates_p, "  (" + role["note"] + ")", size=8.5, colour=MID_GREY, italic=True)

    # Mandate sentence
    if role["mandate"]:
        m_p = doc.add_paragraph()
        spacing(m_p, before=1, after=1)
        normal_run(m_p, role["mandate"], size=9.5)

    # Bullets
    for bullet in role["bullets"]:
        b_p = doc.add_paragraph(style='List Bullet')
        b_p.paragraph_format.left_indent = Inches(0.2)
        spacing(b_p, before=0, after=1)
        normal_run(b_p, bullet, size=9.5)

    # Tech / sector footer
    if role["tech"]:
        t_p = doc.add_paragraph()
        spacing(t_p, before=1, after=1)
        normal_run(t_p, role["tech"], size=8, colour=MID_GREY, italic=True)

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# EARLIER CAREER
# ══════════════════════════════════════════════════════════════════════════════
ec_heading = doc.add_paragraph()
spacing(ec_heading, before=4, after=2)
bold_run(ec_heading, "EARLIER CAREER", size=10, colour=DARK_NAVY)

earlier = [
    ("Charles Stanley & Co.", "Infrastructure Consultant", "~2007"),
    ("Comunica", "Operations Manager", "~2005"),
    ("Ipitomi Ltd", "Co-Founder & MD (subscription managed-services startup)", "~2001"),
    ("Virgin Media", "Non-Standard Technical Design Authority", "1998–2000"),
    ("Whittington Insurance Markets", "Head of Support & Networks", "~1996"),
]

for org, title, period in earlier:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.2)
    spacing(p, before=0, after=1)
    bold_run(p, org, size=9)
    normal_run(p, f"  —  {title}  ({period})", size=9)

ni_p = doc.add_paragraph()
spacing(ni_p, before=3, after=2)
bold_run(ni_p, "Notable: ", size=9)
normal_run(ni_p,
    "Northern Ireland Schools Connectivity Programme (Virgin Media / DoE, 1998–2000) — Non-Standard "
    "Design Authority for broadband connectivity across all 2,000 NI schools; 100% completion under 2-year DoE mandate.",
    size=9)

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# TECHNICAL PROFICIENCIES
# ══════════════════════════════════════════════════════════════════════════════
tp_heading = doc.add_paragraph()
spacing(tp_heading, before=4, after=2)
bold_run(tp_heading, "TECHNICAL PROFICIENCIES", size=10, colour=DARK_NAVY)

tech_items = [
    ("Networking & Infrastructure:",
     "LAN, WAN, Wi-Fi, SD-WAN, SASE, data centre design & transition, low-latency networking, structured cabling"),
    ("Platforms & Cloud:",
     "Microsoft Azure, AWS, Microsoft 365/O365, virtualisation"),
    ("Vendors:",
     "Cisco, Juniper, Microsoft, AWS, BT, Vodafone, Arista, Palo Alto, VMware"),
    ("Tools:",
     "Microsoft Project, Visio, Trello, Oracle ERP, Pipedrive, Google Suite"),
]

for label, content in tech_items:
    p = doc.add_paragraph()
    spacing(p, before=0, after=1)
    bold_run(p, label + "  ", size=9)
    normal_run(p, content, size=9)

add_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CERTIFICATIONS & CLEARANCES (2-column table)
# ══════════════════════════════════════════════════════════════════════════════
cc_heading = doc.add_paragraph()
spacing(cc_heading, before=4, after=2)
bold_run(cc_heading, "CERTIFICATIONS & CLEARANCES", size=10, colour=DARK_NAVY)

certs = [
    ("SC Clearance",       "Active — Government & Military"),
    ("NPPV3 Clearance",    "Active — National Police Vetting"),
    ("PRINCE2",            "Knowledge current (active certification expired)"),
    ("ITIL Fundamentals",  "Knowledge current (active certification expired)"),
    ("CCNA",               "Knowledge current (active certification expired)"),
    ("SASE Fundamentals",  "Knowledge current (active certification expired)"),
    ("H&S / CDM2015",      "H&S Awareness  |  CDM2015 Awareness"),
]

c_tbl = doc.add_table(rows=len(certs), cols=2)
c_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
c_tbl.style = 'Table Grid'

for ri, (label, value) in enumerate(certs):
    c0 = c_tbl.cell(ri, 0)
    c1 = c_tbl.cell(ri, 1)
    for cell, text, bold in [(c0, label, True), (c1, value, False)]:
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.size = Pt(8.5)
        r.bold = bold
        r.font.color.rgb = BODY_BLACK
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if ri % 2 == 0:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EEF2F7')
            tc_pr.append(shd)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ENGAGEMENT MODEL footer note
# ══════════════════════════════════════════════════════════════════════════════
add_rule(doc)
em_p = doc.add_paragraph()
em_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(em_p, before=3, after=0)
normal_run(em_p,
    "Vehicle: Hiloka Limited  |  Target rate: £650/day  |  Availability: within 2 weeks  |  "
    "SoW / output-based contracts  |  UK-based; proven European and US delivery footprint",
    size=8.5, colour=MID_GREY, italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/russellbatchelor/projects/Russell Batchelor CV/team_outputs/CV_Russell_Batchelor_2page.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
